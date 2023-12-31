from aiogram.dispatcher.storage import FSMContext
from aiogram.dispatcher.filters import Command
from aiogram import types
from aiogram.types import CallbackQuery
from aiogram.dispatcher.filters.builtin import CommandStart


from loader import dp
from loader import db
from loader import bot
# from loader import logger

from states.letter import State_list

import io


from keyboards.inline.callback_data import template
from keyboards.default.menu_start import menu_start_admin, menu_start_user
from keyboards.inline.choise_template import choice
from keyboards.inline.choise_cont_user import letter_choise_cont_user
from keyboards.inline.callback_data import template, letter_cont_user_choise

import pymysql

import urllib.request

import docx
import io
import urllib.request
import re
from docx2pdf import convert
import aiohttp

import subprocess


from datetime import datetime

from data import config

import platform

os_name = platform.system()


month_dict = {
    '01': 'января',
    '02': 'февраля',
    '03': 'марта',
    '04': 'апреля',
    '05': 'мая',
    '06': 'июня',
    '07': 'июля',
    '08': 'августа',
    '09': 'сентября',
    '10': 'октября',
    '11': 'ноября',
    '12': 'декабря'
}
current_datetime = datetime.now()

# Extract and format the date
current_date = current_datetime.strftime('%Y-%m-%d')
doc_data = current_datetime.strftime(
    '%d %m %Y')
doc_data = doc_data.split()
doc_data[1] = month_dict[f'{doc_data[1]}']
doc_data = ' '.join(map(str, doc_data))


class Dict_user:
    dict_user_none_text = {}
    dict_user_indx = {}
    dict_user_full_text = {}
    dict_user_answer = {}
    dict_user_call = {}
    dict_column_name = {"fio": f'имя пользователя',
                        "status": f'статус пользователя',
                        "adress": f'адрес отправителя (ваш адрес) и почтовый индекс (если известно)',
                        "phone": f'ваш номер телефона',
                        "email": f'email пользователя',
                        "inn": f'ИНН пользователя',
                        "pasport": f'пасспорт пользователя',
                        "born": f'дату рождения пользователя',
                        "comment": f'комментарий пользователя',
                        #
                        "cont_org": f'наименование организации получателя',
                        "cont_ogrn": f'ОГРН организации получателя',
                        "cont_adress": f'почтовый адрес организации получателя',
                        "cont_phone": f'телефон представителя организации получателя',
                        "cont_email": f'Email организации получателя',
                        "cont_inn": f'ИНН организации получателя',
                        "cont_pasport": f'паспорт получателя',
                        "cont_born": f'дату рождения получателя',
                        "cont_comment": f'комментарий (статус) организации получателя',
                        "cont_fio": f'ФИО руководителя организации получателя',
                        "cont_headstatus": f'должность руководителя организации получателя',
                        "cont_fiocont": f'ФИО представителя организации получателя',
                        "cont_link": f'ссылку на организацию получателя',
                        "doc_text": f'дополнительный текст'}
    dict_doc_text = {}


@dp.callback_query_handler(text='Создать письмо')
async def enter_test(call: CallbackQuery):
    await choice(call.message)


@dp.callback_query_handler(template.filter())
async def letter(call: CallbackQuery, callback_data: dict, state: FSMContext):
    # logger.info(f"нажата кнопка {call}")
    await call.answer(cache_time=60)
    quantity = callback_data.get('id')

    sql = f"SELECT link FROM template WHERE id = '{quantity}'"
    link_template = db.execute(sql, fetchone=True)
    link_template = link_template[0]
    response = urllib.request.urlopen(link_template)
    async with aiohttp.ClientSession() as session:
        async with session.get(link_template) as response:
            if response.status == 200:
                content = await response.read()

                with io.BytesIO(content) as f:
                    doc = docx.Document(f)

                    full_text = []
                    for paragraph in doc.paragraphs:
                        if '{{' in paragraph.text:
                            text = paragraph.text
                            pattern = r'\{\{(.+?)\}\}'
                            matches = re.findall(pattern, text)
                            full_text.extend(matches)

    Dict_user.dict_user_full_text[f'{call.message.chat.id}'] = full_text
    Dict_user.dict_user_call[f'{call.message.chat.id}'] = link_template
    await letter_choise_cont_user(message=call.message, state=state)


@dp.callback_query_handler(letter_cont_user_choise.filter())
async def letter_2(call: CallbackQuery, callback_data: dict, state: FSMContext):
    await call.answer(cache_time=60)
    quantity = callback_data.get('id')
    sql = f"SELECT * FROM cont_users WHERE id='{quantity}'"
    cont_card = db.execute(sql, fetchall=True, commit=True)
    cont_card = cont_card[0]
    sql = f"SELECT * FROM users WHERE id='{call.message.chat.id}'"
    my_card = db.execute(sql, fetchall=True, commit=True)
    my_card = my_card[0]

    dict_oper = {"fio": f'{my_card[1]}',
                 "status": f'{my_card[2]}',
                 "adress": f'{my_card[3]}',
                 "phone": f'{my_card[4]}',
                 "email": f'{my_card[5]}',
                 "inn": f'{my_card[6]}',
                        "pasport": f'{my_card[7]}',
                        "born": f'{my_card[8]}',
                        "comment": f'{my_card[9]}',
                        "cont_org": f'{cont_card[0]}',
                        "cont_ogrn": f'{cont_card[1]}',
                        "cont_adress": f'{cont_card[2]}',
                        "cont_phone": f'{cont_card[3]}',
                        "cont_email": f'{cont_card[4]}',
                        "cont_inn": f'{cont_card[5]}',
                        "cont_pasport": f'{cont_card[6]}',
                        "cont_born": f'{cont_card[7]}',
                        "cont_comment": f'{cont_card[8]}',
                        "cont_fio": f'{cont_card[9]}',
                        "cont_headstatus": f'{cont_card[10]}',
                        "cont_fiocont": f'{cont_card[11]}',
                        "cont_link": f'{cont_card[12]}',
                        "data": doc_data,
                        "doc_text": ''}

    lst_none = []
    for x in Dict_user.dict_user_full_text[f'{call.message.chat.id}']:
        if dict_oper[x] == "":
            lst_none.append(x)
    await state.update_data(
        {
            'dict_oper': dict_oper,
            'lst_none': lst_none,
            'name': quantity
        }
    )
    if len(lst_none) != 0:
        await user_none(message=call.message, state=state)
    else:
        await call.message.answer(text="Письмо создаётся...")
        await letter_3(message=call.message, state=state)


async def letter_3(message: types.Message, state: FSMContext):
    data = await state.get_data()
    dict_oper = data['dict_oper']
    link_template = Dict_user.dict_user_call[f'{message.chat.id}']

    response = urllib.request.urlopen(link_template)
    async with aiohttp.ClientSession() as session:
        async with session.get(link_template) as response:
            if response.status == 200:
                content = await response.read()

                with io.BytesIO(content) as f:
                    doc = docx.Document(f)

                    sql = f"SELECT name FROM users WHERE id='{message.chat.id}'"
                    name = db.execute(
                        sql, fetchall=True, commit=True)[0][0]

                    for oper in Dict_user.dict_user_full_text[f'{message.chat.id}']:
                        for paragraph in doc.paragraphs:
                            text = paragraph.text
                            search_str = f'{{{{{oper}}}}}'
                            paragraph.text = text.replace(
                                search_str, dict_oper[oper])
                            try:
                                paragraph.text = text.replace(
                                    {{{{{'data'}}}}}, doc_data)
                            except:
                                pass

                    dir_name = f"docs"
                    file_name = dir_name + \
                        f"/{name}_{data['name']}_{current_date}.docx".replace(
                            ' ', '_')
                    doc.save(file_name)
                    if os_name == 'Windows':
                        convert(file_name, file_name[:-5] + ".pdf")
                    elif os_name == 'Linux':
                        command = ['libreoffice', '--convert-to',
                                   'pdf', '--outdir', dir_name, file_name]
                        subprocess.run(command, check=True)
                    if str(message.chat.id) in config.ADMINS:
                        await bot.send_document(message.chat.id, open(file_name[:-5] + ".pdf", 'rb'), reply_markup=menu_start_admin)
                    else:
                        await bot.send_document(message.chat.id, open(file_name[:-5] + ".pdf", 'rb'), reply_markup=menu_start_user)


async def user_none(message: types.Message, state: FSMContext):
    data = await state.get_data()
    Dict_user.dict_user_none_text[f'{message.chat.id}'] = data['lst_none']
    Dict_user.dict_user_indx[f"{message.chat.id}"] = 0
    numb_0 = data['lst_none'][0]
    await message.answer(text=f'Введите {Dict_user.dict_column_name[f"{numb_0}"]}')
    await State_list.first()


async def letter_next(call: CallbackQuery, callback_data: dict, state: FSMContext):
    quantity = callback_data.get('name')
    sql = f"SELECT * FROM cont_users WHERE id='{quantity}'"
    cont_card = db.execute(sql, fetchall=True, commit=True)
    cont_card = cont_card[0]
    sql = f"SELECT * FROM users WHERE id='{call.message.chat.id}'"
    my_card = db.execute(sql, fetchall=True, commit=True)
    my_card = my_card[0]

    dict_oper = {"fio": f'{my_card[1]}',
                 "status": f'{my_card[2]}',
                 "adress": f'{my_card[3]}',
                 "phone": f'{my_card[4]}',
                 "email": f'{my_card[5]}',
                 "inn": f'{my_card[6]}',
                        "pasport": f'{my_card[7]}',
                        "born": f'{my_card[8]}',
                        "comment": f'{my_card[9]}',
                        "cont_org": f'{cont_card[0]}',
                        "cont_ogrn": f'{cont_card[1]}',
                        "cont_adress": f'{cont_card[2]}',
                        "cont_phone": f'{cont_card[3]}',
                        "cont_email": f'{cont_card[4]}',
                        "cont_inn": f'{cont_card[5]}',
                        "cont_pasport": f'{cont_card[6]}',
                        "cont_born": f'{cont_card[7]}',
                        "cont_comment": f'{cont_card[8]}',
                        "cont_fio": f'{cont_card[9]}',
                        "cont_headstatus": f'{cont_card[10]}',
                        "cont_fiocont": f'{cont_card[11]}',
                        "cont_link": f'{cont_card[12]}',
                        "data": doc_data,
                        "doc_text": Dict_user.dict_doc_text[f'{call.message.chat.id}']}

    lst_none = []
    print(Dict_user.dict_user_full_text[f'{call.message.chat.id}'])
    for x in Dict_user.dict_user_full_text[f'{call.message.chat.id}']:
        if dict_oper[x] == "":
            lst_none.append(x)
    await state.update_data(
        {
            'dict_oper': dict_oper,
            'lst_none': lst_none,
            'name': quantity
        }
    )
    if len(lst_none) != 0:
        await user_none(message=call.message, state=state)
    else:
        await call.message.answer(text="Письмо создаётся...")
        await letter_3(message=call.message, state=state)


@dp.message_handler(state=State_list.Q0)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    print(name)
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]
    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q1)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]
    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q2)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)
    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q3)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q4)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q5)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q6)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q7)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q8)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q9)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q10)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q11)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q12)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q13)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q14)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q15)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q16)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q17)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q18)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q19)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()


@dp.message_handler(state=State_list.Q20)
async def answer_lst(message: types.Message, state: FSMContext):
    data = await state.get_data()
    name = data['name']
    user_id = Dict_user.dict_user_indx[f"{message.chat.id}"]
    user_text = Dict_user.dict_user_none_text[f'{message.chat.id}']
    oper = user_text[user_id]

    if oper == 'doc_text':
        Dict_user.dict_doc_text[f'{message.chat.id}'] = message.text
    elif 'cont_' in oper:
        oper = oper.replace("cont_", "")
        sql = f"UPDATE cont_users SET {oper}=%s WHERE user=%s AND id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id, name), commit=True)
    else:
        sql = f"UPDATE users SET {oper}=%s WHERE id=%s;"
        db.execute(sql, parameters=(
            message.text, message.chat.id), commit=True)

    try:
        Dict_user.dict_user_indx[f"{message.from_user.id}"] += 1
        await message.answer(f"Напиши {Dict_user.dict_column_name[f'{user_text[user_id+1]}']}")
    except:
        call = CallbackQuery()
        call['message'] = message
        await letter_next(call=call, callback_data={'@': 'letter_choise', 'name': f'{name}'}, state=state)
        await state.finish()
